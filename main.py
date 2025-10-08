import streamlit as st
import os
import nest_asyncio
import asyncio
from dotenv import load_dotenv
import io
from docx import Document

# Import Playwright-based scraper
from scrape_playwright import scrape_website_playwright as scrape_website
from crawler import WebCrawler

# Helper functions for cleaning and parsing
from scrape import extract_body_content, clean_body_content

# LLM cleaning utility
from llm import clean_content_via_llm, chunk_content_via_llm

# Logger
from logger import logger

# Load environment variables
load_dotenv()

# Detect scraping mode
SCRAPING_MODE = os.getenv("SCRAPING_MODE", "playwright")

# ───────────────────────────────────────── UI SETUP ─────────────────────────────────────────
st.set_page_config(page_title="PageMiner_V1.0", layout="wide")
st.title("🕸️ PageMiner V1.0 — Playwright Edition")

st.sidebar.header("Configuration")
st.sidebar.info(f"🧠 Scraping Mode: **{SCRAPING_MODE.title()}**")

# Input: Website URL
url = st.text_input("🔗 Enter Website URL", placeholder="https://www.example.com")

# Sidebar crawl options
st.sidebar.header("Crawling Options")
crawl_mode = st.sidebar.radio("Scraping Mode", ["Single Page", "Crawl Website"])

if crawl_mode == "Crawl Website":
    max_pages = st.sidebar.slider("Max Pages to Crawl", 1, 500, 10)
    delay = st.sidebar.slider("Delay Between Requests (seconds)", 0, 20, 4)
    same_domain_only = st.sidebar.checkbox("Limit to Same Domain", value=True)


# ───────────────────────────────────────── ASYNC CRAWL ─────────────────────────────────────────
async def async_crawl_website(url, max_pages, delay, same_domain_only):
    progress_bar = st.progress(0)
    status = st.empty()
    stats_container = st.container()
    with stats_container:
        col1, col2, col3, col4 = st.columns(4)
        pages_metric = col1.empty()
        links_metric = col2.empty()
        queue_metric = col3.empty()
        current_metric = col4.empty()

    def progress_callback(data):
        progress_bar.progress(data["progress_percentage"] / 100)
        status.text(f"🔄 {data['message']}")
        pages_metric.metric(
            "Pages Scraped", f"{data['visited_count']}/{data['max_pages']}"
        )
        links_metric.metric("Total Links Found", data["total_links_found"])
        queue_metric.metric("Queue Size", data["queue_size"])

        if data["current_url"]:
            display_url = (
                (data["current_url"][:47] + "...")
                if len(data["current_url"]) > 50
                else data["current_url"]
            )
            current_metric.text(f"Current: {display_url}")

    crawler = WebCrawler(
        max_pages=max_pages,
        delay=delay,
        same_domain_only=same_domain_only,
        progress_callback=progress_callback,
    )
    scraped_data = await crawler.crawl_website(url)
    return crawler, scraped_data


# ───────────────────────────────────────── MAIN ACTION ─────────────────────────────────────────

nest_asyncio.apply()


if st.button("🚀 Start Scraping"):
    if not url:
        st.warning("⚠️ Please enter a valid URL.")
    else:
        if crawl_mode == "Single Page":
            st.info("📄 Scraping single page using Playwright...")
            html = asyncio.run(scrape_website(url))
            body = extract_body_content(html)
            # First perform basic cleaning, then enhance with LLM
            raw_cleaned = clean_body_content(body)
            # Store raw (basic) cleaned content for display
            st.session_state.raw_content = raw_cleaned
            try:
                llm_cleaned = clean_content_via_llm(raw_cleaned)
            except Exception as e:
                logger.error(f"LLM cleaning failed: {e}")
                llm_cleaned = raw_cleaned

            # Store both raw and LLM‑cleaned versions
            st.session_state.cleaned_content = llm_cleaned
            # Keep backward‑compatible key used by the UI
            st.session_state.dom_content = llm_cleaned
            st.session_state.scraped_urls = [url]

        else:
            st.info(f"🌐 Crawling up to {max_pages} pages...")
            crawler, scraped_data = asyncio.run(
                async_crawl_website(url, max_pages, delay, same_domain_only)
            )

            combined_content = crawler.get_all_content()
            stats = crawler.get_crawl_stats()

            # Clean the combined content using LLM
            try:
                llm_cleaned = clean_content_via_llm(combined_content)
            except Exception as e:
                logger.error(f"LLM cleaning failed: {e}")
                llm_cleaned = combined_content

            # Store raw and cleaned versions
            st.session_state.raw_content = combined_content
            st.session_state.cleaned_content = llm_cleaned
            st.session_state.dom_content = (
                llm_cleaned  # keep existing key for backward compatibility
            )
            st.session_state.scraped_urls = stats["visited_urls"]
            st.session_state.crawl_stats = stats

            st.success(f"✅ Completed: Scraped {stats['pages_scraped']} pages")

# ───────────────────────────────────────── VIEW CONTENT ─────────────────────────────────────────
if "dom_content" in st.session_state:
    st.subheader("📄 Scraped Content")
    with st.expander("🔍 View LLM‑Cleaned Content"):
        st.text_area("Cleaned Text", st.session_state.dom_content, height=400)
        # Add download button for the cleaned content as a Word document
        doc2 = Document()
        doc2.add_paragraph(st.session_state.dom_content)
        buf2 = io.BytesIO()
        doc2.save(buf2)
        buf2.seek(0)
        st.download_button(
            label="Download Cleaned Content (DOCX)",
            data=buf2,
            file_name="cleaned_content.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        # Button to chunk the cleaned content
        if st.button("🔀 Chunk Cleaned Content"):
            # Use LLM to create meaningful chunks
            st.session_state.cleaned_chunks = chunk_content_via_llm(
                st.session_state.dom_content
            )
            st.success(f"Created {len(st.session_state.cleaned_chunks)} chunks")

    # Show raw (basic) cleaned content if available
    if "raw_content" in st.session_state:
        with st.expander("🔎 View Raw Cleaned Content (pre‑LLM)"):
            st.text_area("Raw Text", st.session_state.raw_content, height=400)
            # Download button for the raw scraped content (pre‑LLM)
            st.markdown("---")
            st.subheader("💾 Download Raw Scraped Data")
            doc_raw = Document()
            doc_raw.add_paragraph(st.session_state.raw_content)
            buf_raw = io.BytesIO()
            doc_raw.save(buf_raw)
            buf_raw.seek(0)
            st.download_button(
                label="Download Raw Content (DOCX)",
                data=buf_raw,
                file_name="raw_content.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    # Display chunks if they exist
    if "cleaned_chunks" in st.session_state:
        st.subheader("📑 Cleaned Content Chunks")
        for idx, chunk in enumerate(st.session_state.cleaned_chunks, 1):
            with st.expander(f"Chunk {idx} (words: {len(chunk.split())})"):
                st.write(chunk)

    if "scraped_urls" in st.session_state:
        st.subheader("🔗 Scraped URLs")
        with st.expander(f"🔗 Scraped URLs ({len(st.session_state.scraped_urls)})"):
            for i, link in enumerate(st.session_state.scraped_urls, 1):
                st.write(f"{i}. {link}")

    if (
        "crawl_stats" in st.session_state
        and st.session_state.crawl_stats["pages_in_queue"] > 0
    ):
        with st.expander(
            f"🕸️ Remaining Queue ({st.session_state.crawl_stats['pages_in_queue']})"
        ):
            for q in st.session_state.crawl_stats["remaining_queue"][:10]:
                st.write(q)
            if len(st.session_state.crawl_stats["remaining_queue"]) > 10:
                st.write("... and more")

# ───────────────────────────────────────── PARSING / LLM ─────────────────────────────────────────
# if "dom_content" in st.session_state:
#     st.subheader("🤖 Ask or Parse the Content")
#     parse_input = st.text_area("Enter your parsing instruction or question")

#     if st.button("🔍 Parse Content"):
#         if parse_input:
#             st.info("🧠 Parsing with LLM...")
#             chunks = split_dom_content(st.session_state.dom_content)
#             result = parse_with_ollama(chunks, parse_input)
#             st.success("✅ Done")
#             st.write(result)
